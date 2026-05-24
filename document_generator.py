"""DOCX document generation for SVEAR Foundation compliance records."""

from datetime import date, datetime
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from config import GENERATED_DOCUMENTS_DIR, SOCIETY_DETAILS


def _safe_folder_name(value: str) -> str:
    """Return a Windows-friendly folder or file segment."""
    cleaned = "".join(ch if ch.isalnum() else "_" for ch in value.strip())
    return "_".join(part for part in cleaned.split("_") if part)


def build_output_path(document_type: str, meeting_type: str, meeting_date: date) -> Path:
    """Build generated_documents/YYYY/Meeting_Type_Date/document.docx."""
    year = str(meeting_date.year)
    folder_name = f"{_safe_folder_name(meeting_type)}_{meeting_date.isoformat()}"
    file_name = f"{_safe_folder_name(document_type)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_dir = GENERATED_DOCUMENTS_DIR / year / folder_name
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / file_name


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def _add_page_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = f"For {SOCIETY_DETAILS['name']}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_heading_block(document: Document, title: str, meeting_date: date, venue: str) -> None:
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(SOCIETY_DETAILS["name"])
    run.bold = True
    run.font.size = Pt(16)

    full_form = document.add_paragraph()
    full_form.alignment = WD_ALIGN_PARAGRAPH.CENTER
    full_form.add_run(SOCIETY_DETAILS["full_form"]).italic = True

    office = document.add_paragraph()
    office.alignment = WD_ALIGN_PARAGRAPH.CENTER
    office.add_run(f"Registered Office: {SOCIETY_DETAILS['registered_office']}")

    document.add_paragraph()
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title.upper())
    title_run.bold = True
    title_run.underline = True
    title_run.font.size = Pt(14)

    meta = document.add_paragraph()
    meta.add_run("Date: ").bold = True
    meta.add_run(meeting_date.strftime("%d %B %Y"))
    meta.add_run("    Place: ").bold = True
    meta.add_run(venue or "New Delhi")


def _add_signature_blocks(document: Document, include_treasurer: bool = False) -> None:
    document.add_paragraph()
    document.add_paragraph("The above record is prepared for signature and physical filing.")
    table = document.add_table(rows=2, cols=3 if include_treasurer else 2)
    table.style = "Table Grid"

    roles = ["President", "Secretary"]
    if include_treasurer:
        roles.append("Treasurer")

    for index, role in enumerate(roles):
        _set_cell_text(table.cell(0, index), "\n\nSignature", bold=False)
        _set_cell_text(table.cell(1, index), role, bold=True)


def _format_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    _add_page_footer(document)


def _add_agenda_list(document: Document, agenda_items: Iterable[str]) -> None:
    document.add_paragraph("Agenda Items:", style=None).runs[0].bold = True
    for item in agenda_items:
        if item.strip():
            document.add_paragraph(item.strip(), style="List Number")


def _add_attendance_table(document: Document, members_present: Iterable[str]) -> None:
    document.add_paragraph("Attendance:", style=None).runs[0].bold = True
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["S. No.", "Name and Designation", "Signature", "Remarks"]
    for col, header in enumerate(headers):
        _set_cell_text(table.cell(0, col), header, bold=True)

    for index, member in enumerate(members_present, start=1):
        cells = table.add_row().cells
        cells[0].text = str(index)
        cells[1].text = member.strip()
        cells[2].text = ""
        cells[3].text = ""


def _normal_notice_text(meeting_type: str, meeting_date: date, meeting_time: str, venue: str) -> str:
    notice_period = "24 hours" if "emergency" in meeting_type.lower() else "15 days"
    return (
        f"Notice is hereby given that a {meeting_type} of {SOCIETY_DETAILS['name']} "
        f"will be held on {meeting_date.strftime('%d %B %Y')} at {meeting_time} "
        f"at {venue}. This notice is issued in accordance with the applicable rules "
        f"and with a notice period of {notice_period}, wherever applicable."
    )


def _resolution_intro(document_type: str, resolution_subject: str) -> str:
    if document_type == "Auditor Appointment Resolution":
        return f"To consider and approve appointment of auditor for {SOCIETY_DETAILS['name']}."
    if document_type == "Bank Account Operation Resolution":
        return "To authorise operation of bank account and related banking transactions."
    if document_type == "Membership Approval Resolution":
        return "To consider and approve admission of members to the society."
    if document_type == "Project Approval Resolution":
        return "To consider and approve the project proposal placed before the meeting."
    return resolution_subject or "To consider and pass the proposed resolution."


def _generate_resolution_text(document_type: str, resolution_subject: str, resolution_text: str) -> str:
    subject = _resolution_intro(document_type, resolution_subject)
    default_resolution = (
        f"RESOLVED THAT the members present hereby approve: {subject}\n\n"
        f"FURTHER RESOLVED THAT the President, Secretary, and Treasurer of "
        f"{SOCIETY_DETAILS['name']} be and are hereby authorised to take all necessary "
        f"steps, sign papers, submit documents, and do all acts required to give effect "
        f"to this resolution."
    )
    return resolution_text.strip() if resolution_text.strip() else default_resolution


def generate_document(
    document_type: str,
    meeting_type: str,
    meeting_date: date,
    meeting_time: str,
    venue: str,
    agenda_items: List[str],
    members_present: List[str],
    chairperson_name: str,
    secretary_name: str,
    resolution_subject: str,
    resolution_text: str,
    output_path: Optional[Path] = None,
) -> Path:
    """Create a professional DOCX document and return the saved file path."""
    if output_path is None:
        output_path = build_output_path(document_type, meeting_type, meeting_date)

    document = Document()
    _format_document(document)
    _add_heading_block(document, document_type, meeting_date, venue)

    lower_type = document_type.lower()

    if "notice" in lower_type:
        document.add_paragraph(_normal_notice_text(meeting_type, meeting_date, meeting_time, venue))
        _add_agenda_list(document, agenda_items)

    elif "agenda" in lower_type:
        document.add_paragraph(
            f"The agenda for the {meeting_type} scheduled on "
            f"{meeting_date.strftime('%d %B %Y')} at {meeting_time} is as follows:"
        )
        _add_agenda_list(document, agenda_items)

    elif "attendance" in lower_type:
        document.add_paragraph(
            f"Attendance sheet for the {meeting_type} held on "
            f"{meeting_date.strftime('%d %B %Y')} at {meeting_time}."
        )
        _add_attendance_table(document, members_present)

    elif "minutes" in lower_type:
        document.add_paragraph(
            f"The {meeting_type} of {SOCIETY_DETAILS['name']} was held on "
            f"{meeting_date.strftime('%d %B %Y')} at {meeting_time} at {venue}."
        )
        document.add_paragraph(f"The meeting was chaired by {chairperson_name or 'the Chairperson'}.")
        document.add_paragraph(f"The proceedings were recorded by {secretary_name or 'the Secretary'}.")
        document.add_paragraph(
            "After confirming that the required quorum was present, the Chairperson "
            "called the meeting to order. The following matters were discussed and recorded:"
        )
        _add_agenda_list(document, agenda_items)
        document.add_paragraph(
            "The members deliberated on the agenda items and authorised the office bearers "
            "to complete necessary compliance, filing, documentation, and follow-up actions."
        )
        _add_attendance_table(document, members_present)

    else:
        document.add_paragraph(
            f"Extract from the proceedings of the {meeting_type} of {SOCIETY_DETAILS['name']} "
            f"held on {meeting_date.strftime('%d %B %Y')} at {meeting_time} at {venue}."
        )
        document.add_paragraph("Subject:", style=None).runs[0].bold = True
        document.add_paragraph(_resolution_intro(document_type, resolution_subject))
        document.add_paragraph("Resolution:", style=None).runs[0].bold = True
        for block in _generate_resolution_text(document_type, resolution_subject, resolution_text).split("\n"):
            if block.strip():
                document.add_paragraph(block.strip())

    include_treasurer = "bank" in lower_type or "audit" in lower_type or "accounts" in lower_type
    _add_signature_blocks(document, include_treasurer=include_treasurer)

    document.save(output_path)
    return output_path
